"""Homework sheets, document export, and submission checking for SHIKSHA."""

from __future__ import annotations

import io
import os
from typing import Any

from google import genai
from google.genai import types

from main import ASSISTANT_NAME, get_api_key, load_env_file

load_env_file()

DEFAULT_TEXT_MODEL = "gemini-2.0-flash"

# Singleton client — reused across calls to avoid "client has been closed" errors
_GENAI_CLIENT: genai.Client | None = None


def get_text_model() -> str:
    text_model = os.getenv("GEMINI_TEXT_MODEL", "").strip()
    if text_model:
        return text_model

    configured_model = os.getenv("GEMINI_MODEL", "").strip()
    if configured_model and "live" not in configured_model.lower():
        return configured_model

    return DEFAULT_TEXT_MODEL


def _client() -> genai.Client:
    """Return the module-level genai.Client singleton, creating it on first call."""
    global _GENAI_CLIENT
    if _GENAI_CLIENT is None:
        _GENAI_CLIENT = genai.Client(api_key=get_api_key())
    return _GENAI_CLIENT


def generate_homework_sheet(details: str, grade_hint: str = "") -> str:
    """Create structured homework content from teacher/parent notes."""
    prompt = f"""You are {ASSISTANT_NAME}, a caring tutor for children aged 4–12.

Write clear, child-friendly homework based on these instructions:
---
{details.strip()}
---
{f"Grade/age: {grade_hint}" if grade_hint else ""}

Format the output as:
# Homework — [date or topic title]

**Subject:** ...
**For:** Student name (leave blank line if unknown)

## Instructions
(short friendly intro)

## Tasks
1. ...
2. ...
3. ...

## Bonus (optional)
...

## Notes for parents
(one short line)

Use simple language. Number every task. Keep it printable."""
    response = _client().models.generate_content(
        model=get_text_model(),
        contents=prompt,
    )
    return (response.text or "").strip()


def check_homework(
    assignment_context: str,
    submission_text: str,
    image_parts: list[types.Part] | None = None,
) -> str:
    """Review a student's homework submission."""
    prompt = f"""You are {ASSISTANT_NAME}, a warm and fair tutor checking a child's homework.

**Assigned homework / what was expected:**
{assignment_context.strip() or "See the submitted work and give general educational feedback."}

**Student submission (text extracted from upload or typed):**
{submission_text.strip() or "(No text — review from image if provided.)"}

Give feedback in this structure:

## Overall
(one encouraging sentence)

## What you did well
- bullet points

## Mistakes / improvements
- gentle corrections with the right answer where needed

## Score (for parents)
X/10 with one line why

## Message for the child
(short, spoken-style encouragement in simple English/Hinglish)

Be kind. Never shame. If submission is empty, ask them to try again."""
    contents: list[Any] = [prompt]
    if image_parts:
        contents.extend(image_parts)

    response = _client().models.generate_content(
        model=get_text_model(),
        contents=contents,
    )
    return (response.text or "").strip()


def extract_upload_text(uploaded_file: Any) -> tuple[str, list[types.Part]]:
    """Return plain text and optional image parts for multimodal check."""
    name = (uploaded_file.name or "").lower()
    raw = uploaded_file.getvalue()
    image_parts: list[types.Part] = []

    if name.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif")):
        mime = uploaded_file.type or "image/jpeg"
        image_parts.append(types.Part.from_bytes(data=raw, mime_type=mime))
        return "", image_parts

    if name.endswith(".txt"):
        return raw.decode("utf-8", errors="replace"), image_parts

    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(raw))
            pages = [p.extract_text() or "" for p in reader.pages]
            return "\n".join(pages).strip(), image_parts
        except Exception as exc:
            return f"[Could not read PDF: {exc}]", image_parts

    if name.endswith(".docx"):
        try:
            from docx import Document

            doc = Document(io.BytesIO(raw))
            return "\n".join(p.text for p in doc.paragraphs if p.text).strip(), image_parts
        except Exception as exc:
            return f"[Could not read Word file: {exc}]", image_parts

    return "[Unsupported file type. Use .txt, .pdf, .docx, or an image.]", image_parts


def extract_upload_text_bytes(
    raw: bytes, filename: str, content_type: str = ""
) -> tuple[str, list[types.Part]]:
    """Same as extract_upload_text but accepts raw bytes + filename — for FastAPI."""
    name = (filename or "").lower()
    image_parts: list[types.Part] = []

    if name.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif")):
        mime = content_type or "image/jpeg"
        image_parts.append(types.Part.from_bytes(data=raw, mime_type=mime))
        return "", image_parts

    if name.endswith(".txt"):
        return raw.decode("utf-8", errors="replace"), image_parts

    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            pages = [p.extract_text() or "" for p in reader.pages]
            return "\n".join(pages).strip(), image_parts
        except Exception as exc:
            return f"[Could not read PDF: {exc}]", image_parts

    if name.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            return "\n".join(p.text for p in doc.paragraphs if p.text).strip(), image_parts
        except Exception as exc:
            return f"[Could not read Word file: {exc}]", image_parts

    return "[Unsupported file type. Use .txt, .pdf, .docx, or an image.]", image_parts


def create_pdf_bytes(title: str, body: str) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, title[:120], new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font("Helvetica", size=11)
    for line in body.replace("\r", "").split("\n"):
        safe = line.encode("latin-1", errors="replace").decode("latin-1")
        pdf.multi_cell(0, 6, safe, new_x="LMARGIN", new_y="NEXT")
    out = pdf.output()
    if isinstance(out, str):
        return out.encode("latin-1")
    return bytes(out)



def create_docx_bytes(title: str, body: str) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=0)
    for line in body.replace("\r", "").split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
